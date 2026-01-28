from docx import Document as DocxDocument
from langchain_core.documents import Document
from ingestion.chunker import chunk_docs
def ingest_docx(path):

    doc = DocxDocument(path)

    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    docs = [
        Document(
            page_content=text,
            metadata={"type": "docx", "source": path}
        )
    ]

    docs = chunk_docs(docs)

    return docs 
